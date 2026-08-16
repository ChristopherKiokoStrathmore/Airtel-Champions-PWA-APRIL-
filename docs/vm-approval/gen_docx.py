#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Build the Airtel Champions App Shadow IT / VM-approval architecture document (.docx).
Compliant with Airtel Africa Shadow IT Policy v2.0 (IT-ShadowIT) and the VM-approval
data-architecture requirements. No en/em dashes are used anywhere (house style)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paths are resolved relative to this script so the document regenerates from any
# checkout. Figures (fig1_asis.png / fig2_tobe.png) live next to this script; the
# .docx is written to the repository root (two levels up from docs/vm-approval).
_HERE = os.path.dirname(os.path.abspath(__file__))
SCR = _HERE
OUT = os.path.normpath(os.path.join(_HERE, "..", "..", "Airtel_Champions_App_ShadowIT_VM_Approval.docx"))

RED   = RGBColor(0xE4, 0x00, 0x2B)
INK   = RGBColor(0x1A, 0x22, 0x33)
BLUE  = RGBColor(0x2C, 0x5A, 0xA0)
GREY  = RGBColor(0x5A, 0x6B, 0x87)
HDRBG = "D9E2F3"
ALTBG = "F2F5FB"
WARNBG= "FDECEC"

doc = Document()

# ---------- base styles ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.08

def style_heading(name, size, color, bold=True):
    s = doc.styles[name]; s.font.name = "Calibri"; s.font.size = Pt(size)
    s.font.bold = bold; s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(12); s.paragraph_format.space_after = Pt(4)
    s.paragraph_format.keep_with_next = True
style_heading("Heading 1", 15, RED)
style_heading("Heading 2", 12.5, INK)
style_heading("Heading 3", 11, BLUE)

# ---------- helpers ----------
def add_field(paragraph, instr):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = instr
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(t); run._r.append(e)

def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexc); tcPr.append(shd)

def set_cell(cell, text, bold=False, size=9.0, color=INK, align="left", bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    for i, line in enumerate(str(text).split("\n")):
        if i > 0: p.add_run().add_break()
        r = p.add_run(line); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
        r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(2)
    if bg: shade(cell, bg)

def add_table(headers, rows, widths=None, header_bg=HDRBG, fs=9.0, zebra=True, hdr_color=INK):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = False
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, size=fs+0.3, color=hdr_color,
                 align="left", bg=header_bg)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        bg = ALTBG if (zebra and ri % 2 == 1) else None
        for ci, val in enumerate(row):
            set_cell(cells[ci], val, size=fs, bg=bg)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def P(text="", size=10.5, bold=False, italic=False, color=INK, align="left",
      after=6, before=0):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "just": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if text:
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.color.rgb = color
    return p

def bullets(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it); r.font.size = Pt(size); r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(3)

def numbers(items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(it); r.font.size = Pt(size); r.font.color.rgb = INK
        p.paragraph_format.space_after = Pt(3)

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)

def figure(path, caption, width=6.2):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    c.paragraph_format.space_after = Pt(8)

def note_box(text):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    set_cell(t.rows[0].cells[0], text, size=9.5, color=INK, bg=WARNBG)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------- page setup (A4) ----------
sec = doc.sections[0]
sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

# footer: Internal + page number
def build_footer(section):
    fp = section.footer.paragraphs[0]; fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Internal   |   Page "); r.font.size = Pt(8.5); r.font.color.rgb = GREY
    add_field(fp, "PAGE")
    r2 = fp.add_run(" of "); r2.font.size = Pt(8.5); r2.font.color.rgb = GREY
    add_field(fp, "NUMPAGES")
    for rr in fp.runs:
        rr.font.size = Pt(8.5); rr.font.color.rgb = GREY
build_footer(sec)

# ============================================================================
# COVER
# ============================================================================
for _ in range(3): doc.add_paragraph()
P("AIRTEL AFRICA", size=13, bold=True, color=RED, align="center", after=2)
P("Shadow IT Deployment Justification and Architecture Document", size=20, bold=True,
  color=INK, align="center", after=4)
P("Virtual Machine / Infrastructure Approval Request", size=13, bold=True, color=BLUE,
  align="center", after=18)
P("Application:  Airtel Champions Field Force and HBB Operations Platform", size=12,
  bold=True, color=INK, align="center", after=2)
P("(Sales Champions, Home Broadband install, Airtel Money agent and ODU retrieval modules)",
  size=10, italic=True, color=GREY, align="center", after=24)

meta = doc.add_table(rows=0, cols=2); meta.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in [
    ("Document classification", "Internal"),
    ("Document version", "1.0 (Draft for IT Architectural Review Board and IT Security)"),
    ("Prepared for", "Airtel Kenya (Airtel Networks Kenya Ltd) - Opco ITD  [confirm Opco]"),
    ("Prepared by", "[Requestor name / business function]"),
    ("Date", "August 2026"),
    ("Governing policy", "Airtel Africa Shadow IT Policy v2.0 (Doc code IT-ShadowIT, released 19-Mar-2026)"),
    ("Mandatory approvers", "CIO, GCIO, GCISO (per policy Section 6)"),
]:
    cells = meta.add_row().cells
    set_cell(cells[0], k, bold=True, size=9.5, bg=HDRBG); set_cell(cells[1], v, size=9.5)
    cells[0].width = Inches(2.3); cells[1].width = Inches(4.0)

doc.add_paragraph()
P("Confidentiality: This document contains architecture and security information about an "
  "Airtel Africa business application. Handle as INTERNAL. Do not distribute outside the "
  "Airtel IT governance, security and approval chain.", size=8.5, italic=True, color=GREY,
  align="center")

doc.add_page_break()

# ============================================================================
# 1. DOCUMENT CONTROL
# ============================================================================
h1("1. Document Control")

h2("1.1 Document Information")
add_table(["Type of information", "Document data"], [
    ["Document title", "Airtel Champions App - Shadow IT Justification and Architecture (VM Approval)"],
    ["Document code", "[to be assigned by Opco ITD]"],
    ["Application / product", "Airtel Champions Field Force and HBB Operations Platform"],
    ["Business function / owner", "[Sales and Distribution / HBB Operations - confirm owner]"],
    ["Requestor", "[Name, title, AAUID]"],
    ["Opco", "Airtel Kenya (Airtel Networks Kenya Ltd)  [confirm]"],
    ["Applicability", "Airtel Africa Plc and its entities"],
    ["Classification", "Internal"],
    ["Version", "1.1"],
    ["Status", "Draft for review and approval"],
], widths=[2.3, 4.0])

h2("1.2 Reviewers (per policy Section 6)")
add_table(["Sl.", "Review body / responsibility", "Name", "Outcome / date"], [
    ["1", "IT Architectural Review Board", "[to be completed]", ""],
    ["2", "IT Security", "[to be completed]", ""],
    ["3", "Opco ITD (sponsoring)", "[to be completed]", ""],
], widths=[0.4, 3.0, 1.8, 1.1])

h2("1.3 Approvers (mandatory)")
P("Policy v2.0 Section 6 states that deployment of a Shadow IT system for a business "
  "process that includes IT is mandatory to be approved by CIO, GCIO and GCISO. This "
  "submission is prepared to seek that approval.", size=10)
add_table(["Sl.", "Role", "Name", "Signatory", "Date"], [
    ["1", "Chief Information Officer (CIO)", "[ ]", "[ ]", "[ ]"],
    ["2", "Group Chief Information Officer (GCIO)", "[ ]", "[ ]", "[ ]"],
    ["3", "Group Chief Information Security Officer (GCISO)", "[ ]", "[ ]", "[ ]"],
], widths=[0.4, 2.9, 1.1, 1.1, 0.9])

h2("1.4 Revision History")
add_table(["Ver.", "Date", "Change description", "Author"], [
    ["1.0", "12 Aug 2026", "Initial submission: as-is architecture, proposed migration to Airtel-managed VMs, VM sizing and controls.", "[Requestor]"],
    ["1.1", "16 Aug 2026", "Updated security controls (Section 10), risk (Section 11) and retirement plan (Section 12) to reflect the interim hardening completed on the current platform: public anon key revoked on all personal-data and credential tables, RLS enforced, and server-side sign in. Corrected an as-is statement: the service-role key is not present in the client bundle.", "[Requestor]"],
], widths=[0.5, 0.9, 3.8, 1.1])

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
h1("Table of Contents")
tp = doc.add_paragraph()
add_field(tp, 'TOC \\o "1-2" \\h \\z \\u')
P("(Right-click the table above in Microsoft Word and choose Update Field to populate page numbers.)",
  size=8.5, italic=True, color=GREY)
doc.add_page_break()

# ============================================================================
# 2. EXECUTIVE SUMMARY
# ============================================================================
h1("2. Executive Summary")
P("The Airtel Champions App is a field force and operations platform used by Airtel sales "
  "champions, Home Broadband (HBB) installers, Airtel Money agents and headquarters teams. "
  "It was built by an external developer and is presently hosted entirely on third-party, "
  "Internet-based Software-as-a-Service (SaaS) platforms (Vercel and Supabase), outside "
  "Airtel Group IT visibility and control. Under Airtel Africa Shadow IT Policy v2.0 this "
  "constitutes Shadow IT.", align="just")
P("This document is the justification and architecture submission required by Section 6 of "
  "that policy. It sets out what the application is, how it is built today, the risks of the "
  "current arrangement, and a proposed to-be architecture that migrates the entire solution "
  "onto Airtel-managed virtual machines inside an Airtel data centre or private cloud, under "
  "Group IT control. It includes the hardware (VM) specification, the sizing basis, the "
  "integration touch points, third-party licences and the security controls, so that the IT "
  "Architectural Review Board, IT Security and the approving CIO / GCIO / GCISO can take a "
  "decision.", align="just")
h3("What is being requested")
bullets([
    "Approval to bring the Airtel Champions App under Group IT management (regularise the Shadow IT).",
    "Provisioning of the virtual machines specified in Section 7 across Production and Non-Production (UAT) environments.",
    "Retirement of the external SaaS subscriptions (Vercel, Supabase, Capgo) once migration is complete, with release of the associated firewall, DNS and licence entries per policy Section 6.",
])

# ============================================================================
# 3. PRODUCT / PROJECT DESCRIPTION
# ============================================================================
h1("3. Product and Project Description")
P("This section answers the VM-approval requirement to state the product or project for "
  "which the hardware is requested.", italic=True, size=9.5, color=GREY)
h2("3.1 Overview")
P("The Airtel Champions App is a Progressive Web Application (PWA) with an optional native "
  "mobile wrapper. It supports four operational domains through a single sign-on screen with "
  "a role and mode selector:")
add_table(["Domain / mode", "Primary users", "Function"], [
    ["Sales (Champions)", "Sales executives, DSEs, zonal managers, HQ", "Field sales activity capture, check-in/out, programmes, reporting and league tables."],
    ["HBB (Home Broadband)", "Installers, supervisors, CX, warehouse, HQ", "Installation job allocation, field capture, supervision and HQ dashboards."],
    ["Airtel Money", "Airtel Money agents and admin", "Agent-facing records and administration."],
    ["ODU Retrieval", "HQ, CX, installers, warehouse", "Recovery of Outdoor Units from ~19,000 inactive HBB customers: upload, call intake, field collection, warehouse validation, reconciliation and payment (KSh 600 per unit)."],
], widths=[1.5, 2.3, 2.5])
h2("3.2 Data held")
bullets([
    "Customer personally identifiable information (PII): names, phone numbers, locations of HBB and inactive customers.",
    "Field staff records and role based access data (phone number plus PIN credentials, session tokens).",
    "Operational records: installations, check-ins, GPS location traces, device photographs and barcodes / IMEIs.",
    "Payment-relevant records for ODU retrieval (payable units, reconciliation batches, KSh amounts, CSV payment exports).",
])
note_box("Because the platform holds customer PII and payment-relevant records, policy Section 6 "
         "requires that it not be operated as Shadow IT without proper controls (encryption, "
         "access control, audit trails and backup per AAISP). Section 10 of this document sets "
         "out how the proposed architecture meets that requirement.")

# ============================================================================
# 4. SHADOW IT JUSTIFICATION
# ============================================================================
h1("4. Shadow IT Justification (Policy Section 6)")
h2("4.1 Business need (what is needed)")
P("Airtel Kenya's sales and HBB operations require a mobile-first tool to direct and measure "
  "a large distributed field force (roughly 5,000 to 6,000 named users across champions, "
  "DSEs, installers, supervisors and agents) and, most recently, to run the ODU retrieval "
  "programme that recovers leased equipment worth material value from about 19,000 inactive "
  "customers. No existing centrally provided Airtel application delivered this combination of "
  "field capture, geo-allocation, barcode capture and reconciliation within the required "
  "timeframe.")
h2("4.2 Why the alternative was used (and why it must now be regularised)")
P("The application was delivered quickly on SaaS platforms to meet an immediate operational "
  "need. That speed came at the cost of central oversight: the solution now runs outside "
  "Airtel IT, on external infrastructure, holding customer PII and payment data. This is "
  "precisely the risk the Shadow IT Policy addresses. The business intent is not to retain an "
  "unsanctioned tool but to bring a proven, in-use application under Group IT control on "
  "Airtel-managed infrastructure.")
h2("4.3 How it will be used")
P("The platform will continue to serve the Sales, HBB, Airtel Money and ODU workflows for "
  "Airtel Kenya field and office staff. Access is role based (see Appendix B). After "
  "migration it will be reached through the standard Airtel ingress (WAF and load balancer), "
  "with all data resident inside the Airtel estate.")
h2("4.4 Operational requirements (scoped in advance, per Section 6)")
add_table(["Item", "Detail"], [
    ["Who maintains the process", "Opco ITD application support, with the incumbent development team transitioned to a managed support contract under IT oversight. Database and OS administration by Opco infrastructure team."],
    ["Access levels", "Role based access (RBAC) enforced in-application; infrastructure access restricted to named administrators via bastion with MFA; least privilege on DB and storage."],
    ["Support model", "Business hours application support with on-call for Severity 1; standard Airtel change, incident and patch management."],
    ["Required duration", "Long term (greater than one year). Per Section 6 this makes a centralised, IT-managed solution mandatory rather than a temporary exception."],
    ["Data handling", "Encryption in transit and at rest, audit trails, backup and password / PIN policy per AAISP (Section 10)."],
], widths=[1.9, 4.4])

# ============================================================================
# 5. EXISTING (AS-IS) ARCHITECTURE
# ============================================================================
h1("5. Existing (As-Is) Architecture")
h2("5.1 Application architecture")
P("The client is a React 18 and TypeScript single-page PWA built with Vite, wrapped for "
  "Android and iOS with Capacitor and updated over the air by Capgo. The browser client talks "
  "directly to a Supabase Cloud backend project. There is no Airtel-hosted middle tier: the "
  "browser holds the API key and calls the backend over HTTPS. Authentication is a custom "
  "phone-number plus PIN scheme with a session token issued by an edge function; Supabase's "
  "own GoTrue auth is present but unused.")
figure(f"{SCR}/fig1_asis.png",
       "Figure 1. Existing (as-is) architecture. The entire solution runs on external SaaS outside Airtel IT control.")

h2("5.2 Network view and components")
P("All components other than the end-user device are external, Internet-reachable SaaS. The "
  "logical components are:")
bullets([
    "Vercel CDN / Edge: static hosting of the compiled PWA (dist/), global CDN delivery.",
    "Supabase Cloud (project xspogpfohjmkykfjadhk): the backend, comprising the components below.",
    "Edge Functions: about 20 Deno-runtime functions using the Hono framework (auth-login, se-login, hbb-*, service-requests, make-server-28f2f653 and others), all currently running with verify_jwt = false.",
    "PostgREST: an automatically generated REST API over PostgreSQL, called by the browser with anon and service-role keys.",
    "PostgreSQL: a single managed database node with Row Level Security (RLS), SECURITY DEFINER stored procedures and pg_cron scheduled jobs.",
    "Realtime: a WebSocket (WSS) channel for live updates.",
    "Supabase Storage: object buckets for ODU documents and photographs.",
    "Capgo: an external over-the-air update service that pushes JavaScript bundles to the installed native app.",
    "Browser-direct third-party endpoints: OpenStreetMap tiles and cdnjs (Leaflet maps), WhatsApp click-to-chat (wa.me).",
])

h2("5.3 Internal integration points")
add_table(["From", "To", "Protocol", "Purpose"], [
    ["Edge Functions", "PostgreSQL / PostgREST", "SQL / REST", "Business logic, RPC calls, reads and writes."],
    ["Browser client", "PostgREST", "HTTPS / REST", "Direct data reads (RLS-scoped)."],
    ["Browser client", "Realtime", "WSS", "Live subscriptions."],
    ["Edge Functions", "Supabase Storage", "HTTPS (S3-style)", "Signed URLs, document and photo objects."],
    ["pg_cron", "PostgreSQL functions", "In-database", "Scheduled sweeps (for example ODU allocation timeout)."],
], widths=[1.5, 1.7, 1.4, 1.7])

h2("5.4 External integration points")
P("System to system (northbound / southbound) and user to system, with protocols:")
add_table(["Integration", "Direction", "Protocol / port", "Notes"], [
    ["User device to Vercel", "User to system (inbound)", "HTTPS / 443", "Loads the PWA."],
    ["User device to Supabase", "User to system (inbound)", "HTTPS / 443, WSS / 443", "Data and realtime; API key in client."],
    ["Capgo to native app", "Southbound (push)", "HTTPS / 443", "OTA JavaScript bundle updates."],
    ["Browser to OpenStreetMap / cdnjs", "Northbound (egress)", "HTTPS / 443", "Map tiles and Leaflet assets."],
    ["Browser to WhatsApp (wa.me)", "Northbound (egress)", "HTTPS / 443", "Click-to-chat deep link."],
    ["Build to Figma", "Design-time only", "HTTPS / 443", "Source of design assets; not a runtime dependency."],
], widths=[1.9, 1.6, 1.4, 1.4])

h2("5.5 Infrastructure architecture and software stack (as-is)")
add_table(["Layer", "As-is technology", "Hosting / ownership"], [
    ["Presentation", "React 18, TypeScript, Vite PWA; Capacitor native shell", "Vercel CDN (external SaaS)"],
    ["Application / API", "Deno runtime, Hono framework (edge functions); PostgREST", "Supabase Cloud (external SaaS)"],
    ["Data", "PostgreSQL (single managed node), pg_cron", "Supabase Cloud (external SaaS)"],
    ["Object storage", "Supabase Storage buckets", "Supabase Cloud (external SaaS)"],
    ["Realtime / messaging", "Supabase Realtime (WebSocket)", "Supabase Cloud (external SaaS)"],
    ["Mobile updates", "Capgo OTA", "Capgo (external SaaS)"],
], widths=[1.4, 2.9, 2.0])
h3("5.6 Operating systems and clustering")
P("Not visible to or controlled by Airtel. The OS, patching, clustering and high availability "
  "of Vercel and Supabase are managed by those providers under their own terms. PostgreSQL is "
  "presented as a single managed node; there is no Airtel-controlled replica or failover.")
h3("5.7 Database")
P("PostgreSQL (managed by Supabase). Security is enforced through RLS and SECURITY DEFINER "
  "RPCs. Row-level identity is not based on a real IdP; the application uses anon and "
  "service-role API keys, with the service-role key capable of bypassing RLS.")
h3("5.8 Middleware")
P("PostgREST (REST layer), the Hono web framework on Deno (edge functions), and Supabase "
  "Realtime. All are provider-operated; Airtel cannot tune, patch or monitor them directly.")
h2("5.9 Proprietary / COTS versus open source (as-is)")
add_table(["Component", "Type", "Licence / commercial basis"], [
    ["Vercel hosting", "Proprietary SaaS", "Paid subscription"],
    ["Supabase Cloud", "Open-source core, delivered as proprietary managed cloud", "Paid subscription"],
    ["Capgo OTA", "Proprietary SaaS", "Paid subscription"],
    ["Figma (design)", "Proprietary SaaS", "Paid subscription (design-time only)"],
    ["PostgreSQL, PostgREST, React, Vite, Radix UI, Leaflet", "Open source", "MIT / Apache / PostgreSQL / BSD (no licence fee)"],
    ["OpenStreetMap tiles", "Open data", "Free (attribution / usage policy applies)"],
], widths=[2.4, 2.1, 1.8])

# ============================================================================
# 6. PROPOSED (TO-BE) ARCHITECTURE
# ============================================================================
h1("6. Proposed (To-Be) Architecture: Airtel-Managed VMs")
h2("6.1 Application architecture")
P("The application is re-platformed onto Airtel-managed virtual machines inside an Airtel "
  "data centre or private cloud. The same application code is retained; only the hosting and "
  "the trust boundary change. End users reach the platform through the Airtel WAF and load "
  "balancer. A web tier serves the static PWA; an application tier runs the API / functions as "
  "containers; a data tier runs PostgreSQL in a highly available primary and replica "
  "configuration with an object store for documents and photographs. Secrets move into a "
  "vault and no API key is exposed to the browser for privileged access.")
figure(f"{SCR}/fig2_tobe.png",
       "Figure 2. Proposed (to-be) architecture on Airtel-managed VMs. All PII and payment data remain inside the Airtel estate.")

h2("6.2 Network view and zoning")
add_table(["Zone", "Components", "Exposure"], [
    ["Ingress", "Airtel WAF plus load balancer (active/active), TLS 1.2+ termination", "Only public entry point"],
    ["Web zone (DMZ)", "Web VM #1 and #2 (Nginx serving PWA static)", "Reachable from LB only"],
    ["Application zone", "API VM #1 and #2 (containerised Node/Deno functions)", "Reachable from web zone and LB"],
    ["Data zone (private subnet)", "PostgreSQL primary and replica, PgBouncer, PostgREST, MinIO object storage", "No inbound Internet; reachable from app zone only"],
    ["Management zone", "Bastion (SSH, MFA), backup and WAL archive, monitoring / SIEM, secrets vault", "Admin access only"],
], widths=[1.7, 3.1, 1.5])

h2("6.3 Internal integration points (to-be)")
add_table(["From", "To", "Protocol / port", "Purpose"], [
    ["Load balancer", "Web VMs", "HTTPS / 443", "Serve PWA static assets."],
    ["Load balancer", "API VMs", "HTTPS / 443", "API and function calls."],
    ["Web VMs", "API VMs", "HTTPS / internal", "Application requests."],
    ["API VMs", "PgBouncer / PostgreSQL", "libpq / 6432, 5432", "Pooled database access."],
    ["API VMs", "MinIO", "S3 API / 9000", "Object storage read and write (server-side encrypted)."],
    ["PostgreSQL primary", "PostgreSQL replica", "Streaming replication / 5432", "High availability and read scaling."],
    ["All nodes", "Vault, monitoring, backup", "HTTPS / agent", "Secrets, telemetry and backup."],
], widths=[1.5, 1.7, 1.5, 1.6])

h2("6.4 External integration points (to-be)")
add_table(["Integration", "Direction", "Protocol / port", "Change from as-is"], [
    ["User to Airtel WAF/LB", "User to system (inbound)", "HTTPS / 443", "Single controlled ingress replaces Vercel and direct Supabase calls."],
    ["Map tiles", "Northbound (egress)", "HTTPS / 443", "Move to a self-hosted or Airtel-sanctioned tile source."],
    ["Airtel PKI / CA", "Internal", "HTTPS", "TLS certificate issuance."],
    ["Airtel AD / SSO (future)", "Internal", "LDAPS / SAML / OIDC", "Optional future move from phone+PIN to enterprise identity."],
    ["Mobile app updates", "Southbound (push)", "HTTPS / 443", "Move OTA in-house or use signed store releases; retire external Capgo."],
], widths=[1.7, 1.6, 1.4, 1.6])

h2("6.5 Infrastructure architecture and software stack (to-be)")
add_table(["Layer", "To-be technology", "Notes"], [
    ["Ingress", "Airtel WAF and load balancer (reuse existing Airtel COTS, for example F5 / Imperva class)", "Active/active, TLS termination"],
    ["Web", "Nginx serving compiled PWA", "Stateless, horizontally scalable"],
    ["Application", "Node.js (or Deno) functions in containers (Docker / Podman)", "Ported from Supabase edge functions"],
    ["API layer", "PostgREST plus PgBouncer", "REST over PostgreSQL, pooled"],
    ["Data", "PostgreSQL 16 with Patroni and etcd", "Primary plus replica, automatic failover"],
    ["Object storage", "MinIO (S3-compatible)", "Server-side encryption; note AGPL licensing (Section 9)"],
    ["Secrets", "HashiCorp Vault", "Removes keys from client and config"],
    ["Observability", "Monitoring / SIEM per Airtel standard", "Central logging and alerting"],
], widths=[1.4, 3.0, 1.9])

h3("6.6 Operating systems and clustering")
bullets([
    "Operating system: Ubuntu Server 22.04 LTS or RHEL 9 per Airtel data-centre standard (confirm with infrastructure team).",
    "Web tier: two or more nodes, active/active behind the load balancer.",
    "Application tier: two or more container hosts, horizontally scalable.",
    "Database tier: PostgreSQL primary plus at least one streaming replica, orchestrated by Patroni with etcd for automatic failover.",
    "Object storage: MinIO single node initially, with the option of distributed (multi-node) mode for redundancy.",
])
h3("6.7 Database")
P("PostgreSQL 16, self-managed on Airtel VMs, retaining the existing RLS policies and "
  "SECURITY DEFINER RPCs. High availability via streaming replication and Patroni. Point-in-"
  "time recovery via continuous WAL archiving to the management zone.")
h3("6.8 Middleware")
P("Nginx (web and reverse proxy), PgBouncer (connection pooling), PostgREST (REST API), the "
  "container runtime (Docker or Podman), Patroni and etcd (database HA), and HashiCorp Vault "
  "(secrets). All are Airtel-operated, patchable and monitorable.")
h2("6.9 Proprietary / COTS versus open source (to-be)")
add_table(["Component", "Type", "Licence / commercial basis"], [
    ["PostgreSQL, Nginx, PostgREST, PgBouncer, Patroni, etcd", "Open source", "PostgreSQL / BSD / MIT / ISC / Apache (no fee)"],
    ["MinIO", "Open source", "AGPLv3 (consider commercial licence to avoid AGPL obligations)"],
    ["HashiCorp Vault", "Open source (source-available)", "BSL 1.1 (confirm licensing / consider OpenBao)"],
    ["Operating system", "Open source or licensed", "Ubuntu LTS (free) or RHEL (paid subscription)"],
    ["WAF and load balancer", "COTS", "Reuse existing Airtel licence"],
    ["Container runtime", "Open source", "Docker CE / Podman (no fee)"],
], widths=[2.6, 1.9, 1.8])

# ============================================================================
# 7. HARDWARE / VM SPECIFICATION AND SIZING
# ============================================================================
h1("7. Hardware / VM Specification and Sizing Justification")
P("This section answers the VM-approval requirement to show how the hardware specification "
  "was reached and to provide the supporting data.", italic=True, size=9.5, color=GREY)

h2("7.1 Sizing basis (supporting data)")
add_table(["Driver", "Observed / estimated value", "Sizing implication"], [
    ["Named users", "~5,000 to 6,000 (champions ~957, DSEs ~2,216, app_users ~2,604, installers ~476, agents and HQ)", "Modest total population."],
    ["Peak concurrency", "Estimated 10 to 15 percent of named users = ~600 to 900 concurrent at campaign peaks", "Two web and two API nodes give headroom and HA."],
    ["Database size", "Current data on the order of a few GB (largest sets: ODU inactive ~19,000 rows, DSE ~2,216, app_users ~2,604)", "32 GB RAM DB nodes cache the working set comfortably; 200 GB data volume allows years of growth."],
    ["Object storage", "ODU device photos and documents: ~19,000 units x ~3 images x ~0.5 MB ~= 28 GB, plus other photos and growth", "Provision 1 TB with room to grow."],
    ["Availability target", "Business-critical field operations and payment records", "Primary plus replica DB, active/active web and app tiers."],
])
note_box("These figures are recommended baselines derived from current data volumes and a "
         "conservative growth allowance. They should be validated by IT through load testing "
         "before final procurement. Values marked [confirm] depend on the target Airtel "
         "environment and Opco standards.")

h2("7.2 Production VM specification")
add_table(["Role", "Qty", "vCPU", "RAM", "Disk", "OS"], [
    ["Web (Nginx + PWA)", "2", "2", "4 GB", "40 GB", "Ubuntu 22.04 LTS / RHEL 9"],
    ["Application / API (containers)", "2", "4", "8 GB", "60 GB", "Ubuntu 22.04 LTS / RHEL 9"],
    ["PostgreSQL primary", "1", "8", "32 GB", "200 GB SSD", "Ubuntu 22.04 LTS / RHEL 9"],
    ["PostgreSQL replica", "1", "8", "32 GB", "200 GB SSD", "Ubuntu 22.04 LTS / RHEL 9"],
    ["MinIO object storage", "1", "4", "8 GB", "1 TB", "Ubuntu 22.04 LTS / RHEL 9"],
    ["Bastion / jump host", "1", "2", "4 GB", "40 GB", "Ubuntu 22.04 LTS / RHEL 9"],
    ["Monitoring + Vault", "1", "4", "8 GB", "100 GB", "Ubuntu 22.04 LTS / RHEL 9"],
], widths=[2.1, 0.5, 0.6, 0.8, 1.0, 1.5], fs=8.7)
P("Production subtotal: 8 VMs, approximately 32 vCPU, 96 GB RAM, and about 1.64 TB of disk "
  "before backup allocation.", size=9.5, italic=True, color=GREY)

h2("7.3 Non-Production (UAT) VM specification")
add_table(["Role", "Qty", "vCPU", "RAM", "Disk", "OS"], [
    ["Web + Application (combined)", "1", "4", "8 GB", "60 GB", "Ubuntu 22.04 LTS / RHEL 9"],
    ["PostgreSQL (single node)", "1", "4", "16 GB", "100 GB SSD", "Ubuntu 22.04 LTS / RHEL 9"],
    ["MinIO object storage", "1", "2", "4 GB", "250 GB", "Ubuntu 22.04 LTS / RHEL 9"],
], widths=[2.1, 0.5, 0.6, 0.8, 1.0, 1.5], fs=8.7)
P("UAT subtotal: 3 VMs, approximately 10 vCPU, 28 GB RAM, 410 GB disk. A scaled-down mirror "
  "for testing and release validation (policy expects code review and post-live support, "
  "addressing the operational instability risk in Section 5 of the policy).", size=9.5,
  italic=True, color=GREY)

h2("7.4 Backup, WAL archive and disaster recovery")
bullets([
    "Backup / WAL archive storage: allow 500 GB to 1 TB in the management zone for database backups and continuous WAL archiving (retention per AAISP).",
    "Disaster recovery: a warm replica at a secondary Airtel site is recommended as a Phase 2 item [confirm DR requirement and RPO/RTO with IT].",
])

h2("7.5 How the specification was derived")
numbers([
    "Start from the named-user population and estimate peak concurrency (10 to 15 percent).",
    "Size the web and application tiers for that concurrency with N+1 redundancy (two nodes each), since the workload is light and mostly I/O bound.",
    "Size the database for the working set to fit in RAM (current data is only a few GB; 32 GB gives large cache headroom) and provision SSD for growth and WAL.",
    "Size object storage from the ODU photo volume estimate plus growth (1 TB).",
    "Add management, bastion, monitoring and secrets nodes for operability and security.",
    "Mirror a reduced footprint for UAT to enable testing and controlled releases.",
])

# ============================================================================
# 8. INTEGRATION TOUCH POINTS (CONSOLIDATED)
# ============================================================================
h1("8. Integration Touch Points (Consolidated)")
add_table(["Source", "Target", "Type", "Protocol / port", "Data", "Auth"], [
    ["User device", "Airtel WAF/LB", "User to system", "HTTPS / 443", "App traffic", "Session token / (future SSO)"],
    ["Web VMs", "API VMs", "System to system", "HTTPS / internal", "API requests", "Internal service auth"],
    ["API VMs", "PostgreSQL", "System to system", "libpq / 5432, 6432", "Business data (PII, payments)", "DB credential (Vault)"],
    ["API VMs", "MinIO", "System to system", "S3 API / 9000", "Documents, photos", "Access key (Vault), SSE"],
    ["Primary DB", "Replica DB", "System to system", "Streaming repl / 5432", "WAL stream", "Replication credential"],
    ["Nodes", "Vault / SIEM / Backup", "System to system", "HTTPS / agent", "Secrets, logs, backups", "mTLS / agent token"],
    ["Browser", "Map tiles (sanctioned)", "Northbound egress", "HTTPS / 443", "Map tiles", "None"],
], widths=[1.2, 1.2, 1.2, 1.3, 1.3, 1.3], fs=8.4)

# ============================================================================
# 9. THIRD-PARTY / PAID LICENCES
# ============================================================================
h1("9. Third-Party and Paid Licences")
h2("9.1 Current (to be retired on migration)")
add_table(["Item", "Type", "Action"], [
    ["Vercel subscription", "Hosting SaaS", "Retire after web tier migration"],
    ["Supabase subscription", "Backend SaaS", "Retire after data and API migration"],
    ["Capgo subscription", "OTA SaaS", "Retire; move updates in-house or to signed store releases"],
    ["Figma subscription", "Design SaaS", "Retain only if still used for design; not a runtime dependency"],
    ["Apple Developer / Google Play", "Mobile store", "Retain if native app is published to stores [confirm]"],
], widths=[2.1, 1.6, 2.6])
h2("9.2 Proposed (new or reused)")
add_table(["Item", "Type", "Basis"], [
    ["Operating system", "OS", "Ubuntu LTS (no fee) or RHEL subscription (paid) [choose per standard]"],
    ["WAF and load balancer", "COTS", "Reuse existing Airtel licence"],
    ["MinIO", "Object storage", "AGPLv3 free, or commercial licence to avoid AGPL obligations [decide]"],
    ["HashiCorp Vault", "Secrets", "BSL free tier, Enterprise optional, or OpenBao (open source) [decide]"],
    ["TLS certificates", "PKI", "Airtel internal CA (no external fee)"],
    ["PostgreSQL, Nginx, PostgREST, PgBouncer, Patroni", "Open source", "No licence fee"],
], widths=[2.1, 1.6, 2.6])

# ============================================================================
# 10. DATA SECURITY AND COMPLIANCE CONTROLS
# ============================================================================
h1("10. Data Security and Compliance Controls (Policy Section 6 and AAISP)")
add_table(["Control area", "As-is position", "To-be control"], [
    ["Data residency", "Customer PII and payment data on external SaaS", "All data inside Airtel data centre / private cloud"],
    ["Encryption in transit", "HTTPS to SaaS", "TLS 1.2+ terminated at Airtel WAF; internal TLS between tiers"],
    ["Encryption at rest", "Provider-managed", "Full-disk / volume encryption; server-side encryption on MinIO; DB-level encryption per standard"],
    ["Identity and access", "Phone plus PIN; universal default PIN weakness; login formerly compared PINs in the browser via the public anon key", "Sign in runs server side and mints short lived authenticated tokens; only the public anon key is client side and it is inert for personal data; RBAC retained and hardened; enforce PIN policy; optional future SSO"],
    ["Row level security", "Previously off or inconsistent on credential tables (public read of plaintext PINs)", "RLS enforced on all data tables; anon revoked; privileged writes only through service-role Edge Functions / SECURITY DEFINER RPCs"],
    ["Audit trails", "Limited central visibility", "Central logging and SIEM; DB audit; admin access via bastion with MFA"],
    ["Backup and recovery", "Provider-managed", "Scheduled backups and WAL archive per AAISP; tested restore; optional DR site"],
    ["Secrets management", "Keys distributed to browser and config", "HashiCorp Vault; no secrets in client or repo"],
    ["Vulnerability and patching", "Not under Airtel control", "Standard Airtel patch and vulnerability management"],
], widths=[1.5, 2.3, 2.5], fs=8.6)
note_box("Current-state security note (updated 16 August 2026). The public read exposure "
         "identified in the privacy review has been closed on the current platform, ahead of the "
         "VM migration. The public anon key has been revoked on every table holding personal "
         "data or credentials (the sales user table, the identity store, the HBB installer, DSE "
         "and supervisor tables, the Airtel Money agent and HQ tables, and the ODU customer "
         "table); row level security is enabled with role scoped policies; and all sign in now "
         "runs server side in Edge Functions that mint short lived authenticated tokens, so the "
         "browser never reads a credential. This was verified from the attacker position: the "
         "anon key returns HTTP 401 on all of these tables, and a zero-PII blind-index login "
         "(with the pepper held outside the database) is in use. Remaining items to close, "
         "tracked in Section 12, are the universal default PIN (1234) with forced change "
         "disabled, restricting authenticated read of the residual plaintext PIN columns at "
         "column level, and nulling those columns once all users are on the blind-index store.")

# ============================================================================
# 11. RISK ASSESSMENT
# ============================================================================
h1("11. Risk Assessment (mapped to Policy Section 5)")
add_table(["Policy risk", "Current exposure", "Mitigation in proposed architecture"], [
    ["Security vulnerabilities", "PII and payment data on external SaaS; no Airtel WAF (the public read of credential tables has been closed: anon revoked, RLS enforced, server-side sign in)", "In-house hosting, WAF, Vault, encryption, hardened auth"],
    ["Compliance violations", "Customer data outside Airtel control and data-residency governance", "Data resident in Airtel estate; AAISP controls; DPA needs removed"],
    ["Reduced visibility and control", "No central monitoring of the stack", "Central SIEM, logging, standard change and patch management"],
    ["Financial implications", "Recurring SaaS subscriptions (Vercel, Supabase, Capgo)", "Retire subscriptions; run on Airtel infrastructure"],
    ["Decreased productivity", "Provider outages outside Airtel remedy", "Airtel-operated HA (replica, active/active tiers)"],
    ["Operational instability", "Limited code review and post-live support", "UAT environment, release process, managed support"],
    ["Data silos and inconsistencies", "Data isolated from Airtel systems", "Path to integrate with Airtel identity and data platforms"],
    ["Integration and scalability", "Not vetted by central IT", "Reviewed architecture; horizontal scaling; standard integration"],
], widths=[1.5, 2.3, 2.5], fs=8.6)

# ============================================================================
# 12. MIGRATION AND SHADOW IT RETIREMENT PLAN
# ============================================================================
h1("12. Migration and Shadow IT Retirement Plan (Policy Section 6)")
add_table(["Phase", "Activities", "Outcome"], [
    ["0. Approval", "IT Architectural Review Board and IT Security review; CIO / GCIO / GCISO approval; VM provisioning", "Sanction and infrastructure ready"],
    ["1. Build", "Stand up VMs, network zones, WAF, DB HA, MinIO, Vault, monitoring; port edge functions to containers", "Airtel-hosted environment ready"],
    ["2. Data migration", "Migrate PostgreSQL schema and data (the ODU infrastructure migration is a required source artefact) and storage objects; validate", "Data resident in Airtel estate"],
    ["3. Security hardening", "Interim on current platform already done: anon revoked on all credential/PII tables, RLS enforced, server-side sign in. Remaining: enforce PIN policy and retire default PIN; column-level restriction and nulling of residual plaintext PINs; move keys to Vault; enable audit and backup", "AAISP-aligned controls"],
    ["4. Cutover", "DNS to Airtel ingress; parallel run; go-live; user comms", "Traffic on Airtel infrastructure"],
    ["5. Retirement", "Decommission Vercel, Supabase and Capgo; release firewall, DNS, domain and licence entries; confirm with IT", "Shadow IT retired per Section 6"],
], widths=[1.3, 3.3, 1.7], fs=8.7)
note_box("Per the v2.0 change in Section 6: while retiring the Shadow IT application, all "
         "entries in Firewall, Domain Server, User Licences and similar are to be released, and "
         "the respective function in IT is to ensure the same.")

# ============================================================================
# 13. ASSUMPTIONS AND ITEMS TO CONFIRM
# ============================================================================
h1("13. Assumptions and Items To Be Confirmed")
bullets([
    "Opco is Airtel Kenya (inferred from Kenyan data: KSh currency, 07xx numbers, Nairobi and Thika locations). Confirm.",
    "Requestor identity, business owner and document code to be completed by the requesting function.",
    "Target VM environment (Airtel private cloud, on-premise virtualisation or sanctioned hyperscaler tenancy) and OS standard (Ubuntu LTS or RHEL).",
    "WAF / load balancer product and existing licence to reuse.",
    "Disaster recovery requirement, RPO and RTO.",
    "MinIO and Vault licensing decisions.",
    "Whether the native mobile app is published to app stores (affects store licences and OTA approach).",
    "Sizing to be validated by load testing before final procurement.",
])

# ============================================================================
# 14. APPROVALS
# ============================================================================
h1("14. Approvals and Sign-off")
P("This submission is presented for review and approval in line with Airtel Africa Shadow IT "
  "Policy v2.0, Section 6.")
add_table(["Role", "Name", "Decision", "Signature", "Date"], [
    ["Requestor / business owner", "", "", "", ""],
    ["IT Architectural Review Board", "", "", "", ""],
    ["IT Security", "", "", "", ""],
    ["Opco ITD", "", "", "", ""],
    ["CIO", "", "", "", ""],
    ["GCIO", "", "", "", ""],
    ["GCISO", "", "", "", ""],
], widths=[2.2, 1.3, 1.0, 1.2, 0.9], fs=8.7)

# ============================================================================
# APPENDICES
# ============================================================================
h1("Appendix A. Abbreviations")
add_table(["Abbreviation", "Meaning"], [
    ["AAISP", "Airtel Africa Information Security Policy"],
    ["API", "Application Programming Interface"],
    ["CDN", "Content Delivery Network"],
    ["CIO / GCIO / GCISO", "Chief Information Officer / Group CIO / Group Chief Information Security Officer"],
    ["COTS", "Commercial Off The Shelf"],
    ["DMZ", "Demilitarised Zone (perimeter network)"],
    ["DSE", "Direct Sales Executive"],
    ["HA", "High Availability"],
    ["HBB", "Home Broadband"],
    ["IMEI", "International Mobile Equipment Identity"],
    ["ODU", "Outdoor Unit (customer premises equipment)"],
    ["Opco / ITD", "Operating Company / IT Department"],
    ["PII", "Personally Identifiable Information"],
    ["PWA", "Progressive Web Application"],
    ["RBAC / RLS", "Role Based Access Control / Row Level Security"],
    ["SaaS", "Software as a Service"],
    ["SSE", "Server Side Encryption"],
    ["WAF", "Web Application Firewall"],
    ["WAL", "Write Ahead Log (PostgreSQL)"],
], widths=[1.8, 4.5])

h1("Appendix B. Role Based Access (RBAC) Summary")
add_table(["System", "Roles"], [
    ["Sales", "developer (master), director, admin, hq_staff, zonal_sales_manager, zonal_business_manager, sales_executive, networks_team"],
    ["HBB", "hbb_installer, hbb_installer_supervisor, hbb_dse, hbb_hq, hbb_cx, hbb_warehouse"],
    ["Airtel Money", "airtel_money_admin, airtel_money_agent"],
], widths=[1.5, 4.8])
P("Credential model: phone number plus PIN, with a session token issued server-side. "
  "Individual login credentials are held separately and are not reproduced in this document.",
  size=9, italic=True, color=GREY)

doc.save(OUT)
print("saved:", OUT)
